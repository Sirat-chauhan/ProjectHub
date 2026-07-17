import os
import re
import json
import pandas as pd
from typing import List, Dict, Any, Tuple
from pypdf import PdfReader
import docx
from bs4 import BeautifulSoup
from sqlalchemy.orm import Session
from sqlalchemy import or_, func

from backend.app.models import Document, DocumentChunk
from backend.app.core.config import settings
from backend.app.core.prompts import get_query_intent_prompt

from openai import OpenAI
import tiktoken


# Initialize OpenAI Client (Make sure OPENAI_API_KEY is in your .env file)
client = OpenAI(api_key=settings.OPENAI_API_KEY)


import pdfplumber

class RAGService:
    def __init__(self):
        # Cache the tokenizer to avoid re-loading on every chunk call
        self._tokenizer = tiktoken.get_encoding("cl100k_base")

    def parse_file(self, file_path: str, file_type: str) -> List[Tuple[str, Dict[str, Any]]]:
        """
        Parses different file formats and returns a list of tuples: (text_content, metadata_dict).
        Each tuple represents a page or segment of the file to support precise citations.
        """
        file_type = file_type.lower()
        segments = []

        if file_type == "pdf":
            try:
                with pdfplumber.open(file_path) as pdf:
                    for idx, page in enumerate(pdf.pages):
                        # Extract standard text trying to preserve horizontal layout
                        text = page.extract_text(layout=True) or page.extract_text() or ""
                        
                        # Extract tabular data if present
                        tables = page.extract_tables()
                        for table in tables:
                            if table:
                                # Clean up None values and join cells with |
                                table_text = "\n".join([" | ".join([str(cell).replace('\n', ' ') if cell else "" for cell in row]) for row in table])
                                text += "\n\n[TABLE DATA]\n" + table_text
                                
                        if text.strip():
                            segments.append((text.strip(), {"page_number": idx + 1}))
            except Exception as e:
                print(f"Error parsing PDF file {file_path}: {str(e)}")

        elif file_type in ["docx", "doc"]:
            try:
                doc = docx.Document(file_path)
                
                all_text_elements = []
                # 1. Extract standard paragraphs
                for p in doc.paragraphs:
                    if p.text.strip():
                        all_text_elements.append(p.text.strip())
                        
                # 2. Extract text hidden inside tables
                for table in doc.tables:
                    for row in table.rows:
                        row_data = []
                        for cell in row.cells:
                            if cell.text.strip():
                                row_data.append(cell.text.strip())
                        if row_data:
                            all_text_elements.append(" | ".join(row_data))

                # Word doesn't have native pages, so we group elements in chunks of 5
                current_paragraph_group = []
                group_counter = 1
                for text_item in all_text_elements:
                    current_paragraph_group.append(text_item)
                    if len(current_paragraph_group) >= 5:
                        text = "\n".join(current_paragraph_group)
                        segments.append((text, {"paragraph_group": group_counter}))
                        current_paragraph_group = []
                        group_counter += 1
                        
                if current_paragraph_group:
                    text = "\n".join(current_paragraph_group)
                    segments.append((text, {"paragraph_group": group_counter}))
            except Exception as e:
                print(f"Error parsing DOCX file {file_path}: {str(e)}")

        elif file_type in ["xlsx", "xls", "csv"]:
            try:
                if file_type == "csv":
                    df = pd.read_csv(file_path)
                    text = df.to_string(index=False)
                    segments.append((text, {"sheet_name": "CSV Data"}))
                else:
                    xls = pd.ExcelFile(file_path)
                    for sheet_name in xls.sheet_names:
                        df = pd.read_excel(xls, sheet_name=sheet_name)
                        if not df.empty:
                            text = df.to_string(index=False)
                            segments.append((text, {"sheet_name": sheet_name}))
            except Exception as e:
                print(f"Error parsing tabular file {file_path}: {str(e)}")

        elif file_type in ["html", "htm"]:
            try:
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    soup = BeautifulSoup(f.read(), "html.parser")
                    for element in soup(["script", "style"]):
                        element.decompose()
                    text = soup.get_text(separator="\n")
                    # Clean up multiple newlines
                    lines = [line.strip() for line in text.split("\n") if line.strip()]
                    clean_text = "\n".join(lines)
                    segments.append((clean_text, {"section": "Full HTML Document"}))
            except Exception as e:
                print(f"Error parsing HTML file {file_path}: {str(e)}")

        else:
            # Fallback to plain text parsing
            try:
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    text = f.read()
                    # Clean up multiple newlines
                    lines = [line.strip() for line in text.split("\n") if line.strip()]
                    clean_text = "\n".join(lines)
                    segments.append((clean_text, {"section": "Full Text Document"}))
            except Exception as e:
                print(f"Error parsing raw text file {file_path}: {str(e)}")

        return segments

    def sentence_aware_chunking(self, text: str, max_tokens: int = 400, overlap_sentences_count: int = 3) -> List[str]:
        """
        Splits a text segment into chunks of complete sentences.
        Ensures no sentences are broken in half, and limits chunks by tokens rather than characters.
        """
        # Use cached tokenizer for OpenAI's text-embedding-3-small
        tokenizer = self._tokenizer

        # Split text into sentences using regex looking for sentence endings OR newlines (to catch list items/headers)
        sentences = re.split(r"(?<=[.!?])\s+|\n+", text)
        sentences = [s.strip() for s in sentences if s.strip()]

        chunks = []
        current_chunk = []
        current_tokens = 0

        for sentence in sentences:
            # Count exact tokens in the sentence
            sentence_tokens = len(tokenizer.encode(sentence))
            
            # If a single sentence exceeds the chunk size, we must yield it alone
            if sentence_tokens >= max_tokens:
                if current_chunk:
                    chunks.append(" ".join(current_chunk))
                    current_chunk = []
                    current_tokens = 0
                chunks.append(sentence)
                continue

            # If adding this sentence exceeds token size, yield the current chunk
            if current_tokens + sentence_tokens > max_tokens:
                chunks.append(" ".join(current_chunk))
                # Form the start of the next chunk using the overlapping sentences
                overlap_sentences = current_chunk[-overlap_sentences_count:] if len(current_chunk) >= overlap_sentences_count else current_chunk
                current_chunk = list(overlap_sentences)
                # Recalculate tokens for the overlap context
                current_tokens = len(tokenizer.encode(" ".join(current_chunk)))

            current_chunk.append(sentence)
            current_tokens += sentence_tokens

        if current_chunk:
            chunks.append(" ".join(current_chunk))

        return chunks

    def embed_text(self, text: str) -> List[float]:
        """
        Generates dense vector embedding using OpenAI API (text-embedding-3-small).
        """
        response = client.embeddings.create(
            input=[text],
            model="text-embedding-3-small"
        )
        return response.data[0].embedding
    def ingest_document(self, db: Session, document_id: int) -> bool:
        """
        Performs the complete document ingestion pipeline:
        1. Reads document record.
        2. Obtains local file path (downloads from Supabase to a temporary local file if in Cloud Mode).
        3. Parses the file page/segment-wise.
        4. Chunks segments using Token-Based Sentence-Aware chunker.
        5. Generates local vector embedding for each chunk.
        6. Stores chunks and vector embeddings in the database.
        7. Clean up the temporary local file on completion.
        """
        doc = db.query(Document).filter(Document.id == document_id).first()
        if not doc:
            return False

        # Obtain local path (downloads from Supabase to a temporary file if in Cloud Mode)
        from backend.app.services.storage import storage_service
        try:
            physical_path = storage_service.get_local_path(doc.file_path)
        except Exception as e:
            print(f"Failed to obtain local file path for RAG parser: {str(e)}")
            return False

        try:
            # Parse text into page/segment units
            segments = self.parse_file(physical_path, doc.file_type)
            if not segments:
                return False

            chunk_index = 0
            db_chunks = []

            for text_content, segment_meta in segments:
                # Chunk the segment text into token-safe full sentences
                text_chunks = self.sentence_aware_chunking(text_content)
                
                for chunk_content in text_chunks:
                    # Embed chunk content
                    vector = self.embed_text(chunk_content)

                    # Prepare SQLAlchemy chunk record
                    db_chunk = DocumentChunk(
                        document_id=doc.id,
                        content=chunk_content,
                        chunk_index=chunk_index,
                        document_name=doc.name,
                        embedding=vector,
                        metadata_json=segment_meta
                    )
                    db_chunks.append(db_chunk)
                    chunk_index += 1

            if db_chunks:
                db.add_all(db_chunks)
                db.flush()  # Let the caller control the final commit (atomic transaction)
                return True
            return False
        finally:
            # Clean up temp file if we are in Cloud Storage mode
            if storage_service.use_supabase:
                try:
                    if os.path.exists(physical_path):
                        os.remove(physical_path)
                except Exception as clean_err:
                    print(f"Failed to clean up temp file {physical_path}: {str(clean_err)}")
    def analyze_query_intent(self, query_text: str) -> Dict[str, Any]:
        """
        Uses OpenAI to classify the query intent. Detects if it's a summary/overview request
        and matches target document categories dynamically using LLM semantic understanding
        to avoid hardcoding rules or keyword lists.
        """
        try:
            prompt = get_query_intent_prompt(query_text)
            
            response = client.chat.completions.create(
                model=settings.OPENAI_MODEL,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.0,
                max_tokens=50,
                response_format={"type": "json_object"}
            )
            
            result = json.loads(response.choices[0].message.content)
            category_val = str(result.get("category", "all")).lower().strip()
            if category_val not in ["client", "team"]:
                category_val = "all"

            return {
                "is_summary": bool(result.get("is_summary", False)),
                "category": category_val
            }
        except Exception as e:
            # Safe fallback to standard similarity search on error
            print(f"[Query Analysis Error] {str(e)}")
            return {"is_summary": False, "category": "all"}

    def query_project_chunks(self, db: Session, project_id: int, query_text: str, milestone_id: int = None, category: str = None, top_k: int = 5) -> List[Dict[str, Any]]:
        """
        Performs pgvector Cosine Distance similarity search scoped by project_id and optionally milestone_id.
        Detects summary/overview queries using LLM intent classification, switching to sequential chunk retrieval.
        Filters out matches below a Cosine Similarity score of 0.15 for standard similarity search.
        """
        # Analyze intent using OpenAI to avoid hardcoding keywords/rules
        intent = self.analyze_query_intent(query_text)

        if intent["is_summary"]:
            # Determine target category (prioritize UI-selected category, fallback to LLM intent category)
            current_category = category
            if not current_category or current_category.strip() == "":
                if intent["category"] in ["client", "team"]:
                    current_category = intent["category"]

            # Query chunks sequentially to preserve the logical reading order
            query = db.query(DocumentChunk).join(
                Document, DocumentChunk.document_id == Document.id
            ).filter(
                Document.project_id == project_id
            )

            if milestone_id is not None:
                query = query.filter(Document.milestone_id == milestone_id)

            if current_category is not None and current_category.strip() != "":
                query = query.filter(or_(Document.category == current_category, Document.category == 'global'))

            # Retrieve sequential document parts (increased limit to cover multi-document projects)
            results = query.order_by(
                DocumentChunk.document_id,
                DocumentChunk.chunk_index
            ).limit(40).all()

            retrieved_chunks = []
            for chunk in results:
                page_info = None
                if chunk.metadata_json:
                    if "page_number" in chunk.metadata_json:
                        page_info = f"Page {chunk.metadata_json['page_number']}"
                    elif "paragraph_group" in chunk.metadata_json:
                        page_info = f"Section {chunk.metadata_json['paragraph_group']}"
                    elif "sheet_name" in chunk.metadata_json:
                        page_info = f"Sheet {chunk.metadata_json['sheet_name']}"
                    elif "section" in chunk.metadata_json:
                        page_info = str(chunk.metadata_json['section'])
                    elif "line_group" in chunk.metadata_json:
                        page_info = f"Lines {chunk.metadata_json['line_group']}"

                retrieved_chunks.append({
                    "document": chunk.document_name,
                    "chunk_index": chunk.chunk_index,
                    "page": page_info,
                    "snippet": chunk.content,
                    "score": 1.0  # Perfect score for direct retrieval to bypass threshold checks
                })
            return retrieved_chunks

        # --- IMPROVEMENT 2: HYBRID SEARCH (Semantic + Keyword via RRF) ---
        # 1. Semantic Search (pgvector)
        query_vector = self.embed_text(query_text)
        distance_expr = DocumentChunk.embedding.cosine_distance(query_vector)

        semantic_query = db.query(
            DocumentChunk,
            distance_expr.label("distance")
        ).join(
            Document, DocumentChunk.document_id == Document.id
        ).filter(
            Document.project_id == project_id
        )

        if milestone_id is not None:
            semantic_query = semantic_query.filter(Document.milestone_id == milestone_id)
        if category is not None and category.strip() != "":
            semantic_query = semantic_query.filter(or_(Document.category == category, Document.category == 'global'))

        semantic_results = semantic_query.order_by(distance_expr).limit(20).all()

        # 2. Keyword Search (PostgreSQL Full-Text Search)
        keyword_query = db.query(DocumentChunk).join(
            Document, DocumentChunk.document_id == Document.id
        ).filter(
            Document.project_id == project_id,
            func.to_tsvector('english', DocumentChunk.content).op('@@')(func.plainto_tsquery('english', query_text))
        )

        if milestone_id is not None:
            keyword_query = keyword_query.filter(Document.milestone_id == milestone_id)
        if category is not None and category.strip() != "":
            keyword_query = keyword_query.filter(or_(Document.category == category, Document.category == 'global'))

        keyword_results = keyword_query.limit(20).all()

        # 3. Reciprocal Rank Fusion (RRF) to combine Semantic and Keyword scores
        k = 60
        rrf_scores = {}
        chunk_map = {}

        # Add semantic scores
        for rank, (chunk, distance) in enumerate(semantic_results):
            score = 1.0 - distance
            if score >= 0.15:  # Apply original threshold
                rrf_scores[chunk.id] = rrf_scores.get(chunk.id, 0.0) + (1.0 / (k + rank + 1))
                chunk_map[chunk.id] = chunk

        # Add keyword scores
        for rank, chunk in enumerate(keyword_results):
            rrf_scores[chunk.id] = rrf_scores.get(chunk.id, 0.0) + (1.0 / (k + rank + 1))
            chunk_map[chunk.id] = chunk

        # Get top_k combined results
        sorted_chunk_ids = sorted(rrf_scores.keys(), key=lambda x: rrf_scores[x], reverse=True)[:top_k]

        retrieved_chunks = []
        for cid in sorted_chunk_ids:
            chunk = chunk_map[cid]
            page_info = None
            if chunk.metadata_json:
                if "page_number" in chunk.metadata_json:
                    page_info = f"Page {chunk.metadata_json['page_number']}"
                elif "paragraph_group" in chunk.metadata_json:
                    page_info = f"Section {chunk.metadata_json['paragraph_group']}"
                elif "sheet_name" in chunk.metadata_json:
                    page_info = f"Sheet {chunk.metadata_json['sheet_name']}"
                elif "section" in chunk.metadata_json:
                    page_info = str(chunk.metadata_json['section'])
                elif "line_group" in chunk.metadata_json:
                    page_info = f"Lines {chunk.metadata_json['line_group']}"

            retrieved_chunks.append({
                "document": chunk.document_name,
                "chunk_index": chunk.chunk_index,
                "page": page_info,
                "snippet": chunk.content,
                "score": float(rrf_scores[cid]) # Note: RRF score is not a 0-1 percentage, but a fusion rank
            })

        return retrieved_chunks


rag_service = RAGService()
